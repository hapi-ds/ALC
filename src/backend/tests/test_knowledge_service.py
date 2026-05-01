"""Unit tests for KnowledgeService.

Tests text extraction (digital PDF, DOCX, plain text), scanned PDF detection,
text chunking, embedding generation, and search result filtering (ABAC + CSV).

References:
    - Task 12.11: Write unit tests for text extraction, chunking,
      embedding generation, and search result filtering
    - Requirement 13: Document Indexing and Vector Embedding
    - Requirement 14: Semantic and Hybrid Search
"""

import io
from unittest.mock import patch

import fitz  # PyMuPDF
import pytest
from docx import Document as DocxDocument

from alcoabase.services.knowledge_service import KnowledgeService, SearchResult


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def knowledge_service() -> KnowledgeService:
    """Create a KnowledgeService instance with default settings."""
    with patch("alcoabase.services.knowledge_service.get_settings") as mock_settings:
        settings = mock_settings.return_value
        settings.model_embedding_dimension = 1024
        settings.opensearch_url = "http://localhost:9200"
        service = KnowledgeService()
    return service


@pytest.fixture
def sample_pdf_bytes() -> bytes:
    """Create a simple digital PDF with extractable text."""
    doc = fitz.open()
    page = doc.new_page()
    text_point = fitz.Point(72, 72)
    page.insert_text(text_point, "This is a test document with sample content.")
    page.insert_text(fitz.Point(72, 100), "It contains multiple lines of text.")
    page.insert_text(fitz.Point(72, 128), "Used for testing text extraction.")
    pdf_bytes = doc.tobytes()
    doc.close()
    return pdf_bytes


@pytest.fixture
def scanned_pdf_bytes() -> bytes:
    """Create a PDF with no extractable text (simulating a scanned document)."""
    doc = fitz.open()
    # Add a page with only an image (no text)
    page = doc.new_page()
    # Create a small RGB pixmap (10x10) to simulate a scanned page
    img = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 10, 10), 0)
    img.clear_with(200)  # Fill with gray
    page.insert_image(fitz.Rect(72, 72, 200, 200), pixmap=img)
    pdf_bytes = doc.tobytes()
    doc.close()
    return pdf_bytes


@pytest.fixture
def sample_docx_bytes() -> bytes:
    """Create a simple DOCX file with text content."""
    doc = DocxDocument()
    doc.add_paragraph("This is a test DOCX document.")
    doc.add_paragraph("It has multiple paragraphs for testing.")
    doc.add_paragraph("Third paragraph with more content.")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# Test: Text Extraction (Task 12.1)
# ---------------------------------------------------------------------------


class TestExtractText:
    """Tests for KnowledgeService.extract_text()."""

    def test_extract_pdf_text(
        self, knowledge_service: KnowledgeService, sample_pdf_bytes: bytes
    ) -> None:
        """extract_text() extracts text from a digital PDF."""
        text = knowledge_service.extract_text(sample_pdf_bytes, "application/pdf")

        assert "test document" in text.lower()
        assert "sample content" in text.lower()
        assert len(text) > 0

    def test_extract_docx_text(
        self, knowledge_service: KnowledgeService, sample_docx_bytes: bytes
    ) -> None:
        """extract_text() extracts text from a DOCX file."""
        content_type = (
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document"
        )
        text = knowledge_service.extract_text(sample_docx_bytes, content_type)

        assert "test DOCX document" in text
        assert "multiple paragraphs" in text

    def test_extract_plain_text(self, knowledge_service: KnowledgeService) -> None:
        """extract_text() decodes plain text from bytes."""
        content = "Hello, this is plain text content.\nWith multiple lines."
        file_bytes = content.encode("utf-8")

        text = knowledge_service.extract_text(file_bytes, "text/plain")

        assert text == content

    def test_extract_text_unsupported_type_raises(
        self, knowledge_service: KnowledgeService
    ) -> None:
        """extract_text() raises ValueError for unsupported content types."""
        with pytest.raises(ValueError, match="Unsupported content type"):
            knowledge_service.extract_text(b"data", "application/zip")

    def test_extract_text_handles_utf8_with_bom(
        self, knowledge_service: KnowledgeService
    ) -> None:
        """extract_text() handles UTF-8 text with BOM."""
        content = "\ufeffHello UTF-8 with BOM"
        file_bytes = content.encode("utf-8")

        text = knowledge_service.extract_text(file_bytes, "text/plain")

        assert "Hello UTF-8 with BOM" in text


# ---------------------------------------------------------------------------
# Test: Scanned PDF Detection (Task 12.2)
# ---------------------------------------------------------------------------


class TestScannedPdfDetection:
    """Tests for KnowledgeService.is_scanned_pdf()."""

    def test_digital_pdf_not_scanned(
        self, knowledge_service: KnowledgeService, sample_pdf_bytes: bytes
    ) -> None:
        """is_scanned_pdf() returns False for PDFs with extractable text."""
        assert knowledge_service.is_scanned_pdf(sample_pdf_bytes) is False

    def test_scanned_pdf_detected(
        self, knowledge_service: KnowledgeService, scanned_pdf_bytes: bytes
    ) -> None:
        """is_scanned_pdf() returns True for PDFs without extractable text."""
        assert knowledge_service.is_scanned_pdf(scanned_pdf_bytes) is True

    def test_ocr_fallback_returns_placeholder(
        self, knowledge_service: KnowledgeService, scanned_pdf_bytes: bytes
    ) -> None:
        """extract_text_with_ocr_fallback() returns placeholder for scanned PDFs."""
        text = knowledge_service.extract_text_with_ocr_fallback(
            scanned_pdf_bytes, "application/pdf"
        )

        assert "OCR_PENDING" in text

    def test_ocr_fallback_extracts_digital_pdf(
        self, knowledge_service: KnowledgeService, sample_pdf_bytes: bytes
    ) -> None:
        """extract_text_with_ocr_fallback() extracts text from digital PDFs normally."""
        text = knowledge_service.extract_text_with_ocr_fallback(
            sample_pdf_bytes, "application/pdf"
        )

        assert "test document" in text.lower()


# ---------------------------------------------------------------------------
# Test: Text Chunking (Task 12.3)
# ---------------------------------------------------------------------------


class TestChunkText:
    """Tests for KnowledgeService.chunk_text()."""

    def test_short_text_single_chunk(
        self, knowledge_service: KnowledgeService
    ) -> None:
        """chunk_text() returns single chunk for text shorter than chunk_size."""
        text = "This is a short text with only a few words."
        chunks = knowledge_service.chunk_text(text, chunk_size=512, overlap=50)

        assert len(chunks) == 1
        assert chunks[0] == text.strip()

    def test_long_text_multiple_chunks(
        self, knowledge_service: KnowledgeService
    ) -> None:
        """chunk_text() splits long text into multiple overlapping chunks."""
        # Create text with 1000 words
        words = [f"word{i}" for i in range(1000)]
        text = " ".join(words)

        chunks = knowledge_service.chunk_text(text, chunk_size=512, overlap=50)

        assert len(chunks) > 1

    def test_chunk_overlap(self, knowledge_service: KnowledgeService) -> None:
        """chunk_text() produces overlapping chunks."""
        # Create text with exactly 600 words
        words = [f"word{i}" for i in range(600)]
        text = " ".join(words)

        chunks = knowledge_service.chunk_text(text, chunk_size=512, overlap=50)

        assert len(chunks) == 2
        # The second chunk should start 462 words in (512 - 50)
        # So it should contain word462 through word599
        assert "word462" in chunks[1]

    def test_empty_text_returns_empty_list(
        self, knowledge_service: KnowledgeService
    ) -> None:
        """chunk_text() returns empty list for empty text."""
        assert knowledge_service.chunk_text("") == []
        assert knowledge_service.chunk_text("   ") == []

    def test_overlap_must_be_less_than_chunk_size(
        self, knowledge_service: KnowledgeService
    ) -> None:
        """chunk_text() raises ValueError if overlap >= chunk_size."""
        with pytest.raises(ValueError, match="Overlap"):
            knowledge_service.chunk_text("some text", chunk_size=50, overlap=50)

        with pytest.raises(ValueError, match="Overlap"):
            knowledge_service.chunk_text("some text", chunk_size=50, overlap=100)

    def test_custom_chunk_size_and_overlap(
        self, knowledge_service: KnowledgeService
    ) -> None:
        """chunk_text() respects custom chunk_size and overlap parameters."""
        words = [f"w{i}" for i in range(100)]
        text = " ".join(words)

        chunks = knowledge_service.chunk_text(text, chunk_size=30, overlap=10)

        # Each chunk should have approximately 30 tokens
        for chunk in chunks:
            token_count = len(chunk.split())
            assert token_count <= 30


# ---------------------------------------------------------------------------
# Test: Embedding Generation (Task 12.4)
# ---------------------------------------------------------------------------


class TestGenerateEmbeddings:
    """Tests for KnowledgeService.generate_embeddings()."""

    def test_returns_correct_number_of_embeddings(
        self, knowledge_service: KnowledgeService
    ) -> None:
        """generate_embeddings() returns one embedding per chunk."""
        chunks = ["chunk one", "chunk two", "chunk three"]
        embeddings = knowledge_service.generate_embeddings(chunks)

        assert len(embeddings) == 3

    def test_embedding_dimension_matches_config(
        self, knowledge_service: KnowledgeService
    ) -> None:
        """generate_embeddings() returns vectors of configured dimension."""
        chunks = ["test chunk"]
        embeddings = knowledge_service.generate_embeddings(chunks)

        assert len(embeddings[0]) == 1024  # model_embedding_dimension

    def test_embeddings_are_normalized(
        self, knowledge_service: KnowledgeService
    ) -> None:
        """generate_embeddings() returns approximately unit-length vectors."""
        chunks = ["test chunk for normalization"]
        embeddings = knowledge_service.generate_embeddings(chunks)

        magnitude = sum(v * v for v in embeddings[0]) ** 0.5
        assert abs(magnitude - 1.0) < 0.01

    def test_empty_chunks_returns_empty_list(
        self, knowledge_service: KnowledgeService
    ) -> None:
        """generate_embeddings() returns empty list for empty input."""
        embeddings = knowledge_service.generate_embeddings([])
        assert embeddings == []

    def test_different_chunks_produce_different_embeddings(
        self, knowledge_service: KnowledgeService
    ) -> None:
        """generate_embeddings() produces different vectors for different chunks."""
        chunks = ["first chunk", "second chunk"]
        embeddings = knowledge_service.generate_embeddings(chunks)

        # Random vectors should be different (extremely unlikely to be equal)
        assert embeddings[0] != embeddings[1]


# ---------------------------------------------------------------------------
# Test: OpenSearch Indexing (Task 12.5)
# ---------------------------------------------------------------------------


class TestIndexDocument:
    """Tests for KnowledgeService.index_document()."""

    def test_index_stores_document(
        self, knowledge_service: KnowledgeService
    ) -> None:
        """index_document() stores document in the in-memory index."""
        chunks = ["chunk 1", "chunk 2"]
        embeddings = [[0.1] * 1024, [0.2] * 1024]

        knowledge_service.index_document(
            document_uuid="2025-00001",
            version="1.0",
            chunks=chunks,
            embeddings=embeddings,
            metadata={"title": "Test Doc"},
        )

        assert "2025-00001:1.0" in knowledge_service._index
        indexed = knowledge_service._index["2025-00001:1.0"]
        assert indexed.document_uuid == "2025-00001"
        assert indexed.chunks == chunks
        assert indexed.metadata["title"] == "Test Doc"

    def test_index_multiple_versions(
        self, knowledge_service: KnowledgeService
    ) -> None:
        """index_document() stores multiple versions separately."""
        knowledge_service.index_document(
            document_uuid="2025-00001",
            version="1.0",
            chunks=["v1 chunk"],
            embeddings=[[0.1] * 1024],
        )
        knowledge_service.index_document(
            document_uuid="2025-00001",
            version="2.0",
            chunks=["v2 chunk"],
            embeddings=[[0.2] * 1024],
        )

        assert "2025-00001:1.0" in knowledge_service._index
        assert "2025-00001:2.0" in knowledge_service._index


# ---------------------------------------------------------------------------
# Test: Hybrid Search + Filtering (Tasks 12.7, 12.8, 12.9)
# ---------------------------------------------------------------------------


class TestHybridSearch:
    """Tests for KnowledgeService.hybrid_search()."""

    def test_search_finds_matching_documents(
        self, knowledge_service: KnowledgeService
    ) -> None:
        """hybrid_search() returns documents matching the query."""
        knowledge_service.index_document(
            document_uuid="2025-00001",
            version="1.0",
            chunks=["This document discusses safety procedures"],
            embeddings=[[0.1] * 1024],
            metadata={"title": "Safety SOP"},
        )

        results = knowledge_service.hybrid_search("safety", user_id=1)

        assert len(results) == 1
        assert results[0].document_uuid == "2025-00001"
        assert results[0].title == "Safety SOP"

    def test_search_returns_empty_for_no_match(
        self, knowledge_service: KnowledgeService
    ) -> None:
        """hybrid_search() returns empty list when no documents match."""
        knowledge_service.index_document(
            document_uuid="2025-00001",
            version="1.0",
            chunks=["This is about chemistry"],
            embeddings=[[0.1] * 1024],
            metadata={"title": "Chemistry Doc"},
        )

        results = knowledge_service.hybrid_search("physics", user_id=1)

        assert len(results) == 0

    def test_search_respects_limit(
        self, knowledge_service: KnowledgeService
    ) -> None:
        """hybrid_search() respects the limit parameter."""
        for i in range(10):
            knowledge_service.index_document(
                document_uuid=f"2025-{i:05d}",
                version="1.0",
                chunks=[f"Document {i} about testing procedures"],
                embeddings=[[0.1] * 1024],
                metadata={"title": f"Doc {i}"},
            )

        results = knowledge_service.hybrid_search("testing", user_id=1, limit=3)

        assert len(results) <= 3

    def test_search_excludes_csv_validation_records(
        self, knowledge_service: KnowledgeService
    ) -> None:
        """hybrid_search() excludes documents with is_csv_validation_record=True."""
        # Index a normal document
        knowledge_service.index_document(
            document_uuid="2025-00001",
            version="1.0",
            chunks=["Normal document about testing"],
            embeddings=[[0.1] * 1024],
            metadata={"title": "Normal Doc", "is_csv_validation_record": False},
        )

        # Index a CSV validation record
        knowledge_service.index_document(
            document_uuid="2025-00002",
            version="1.0",
            chunks=["CSV validation record about testing"],
            embeddings=[[0.1] * 1024],
            metadata={"title": "CSV Record", "is_csv_validation_record": True},
        )

        results = knowledge_service.hybrid_search("testing", user_id=1)

        # Only the normal document should be returned
        assert len(results) == 1
        assert results[0].document_uuid == "2025-00001"

    def test_search_abac_filtering_excludes_unauthorized(
        self, knowledge_service: KnowledgeService
    ) -> None:
        """hybrid_search() excludes documents user lacks permission for."""
        # Document accessible to user 1 only
        knowledge_service.index_document(
            document_uuid="2025-00001",
            version="1.0",
            chunks=["Restricted document about protocols"],
            embeddings=[[0.1] * 1024],
            metadata={
                "title": "Restricted Doc",
                "permitted_user_ids": {1, 3},
            },
        )

        # Document accessible to all
        knowledge_service.index_document(
            document_uuid="2025-00002",
            version="1.0",
            chunks=["Public document about protocols"],
            embeddings=[[0.1] * 1024],
            metadata={"title": "Public Doc"},
        )

        # User 1 should see both
        results_user1 = knowledge_service.hybrid_search("protocols", user_id=1)
        assert len(results_user1) == 2

        # User 2 should only see the public document
        results_user2 = knowledge_service.hybrid_search("protocols", user_id=2)
        assert len(results_user2) == 1
        assert results_user2[0].document_uuid == "2025-00002"

    def test_search_abac_none_permits_all_users(
        self, knowledge_service: KnowledgeService
    ) -> None:
        """hybrid_search() permits all users when permitted_user_ids is None."""
        knowledge_service.index_document(
            document_uuid="2025-00001",
            version="1.0",
            chunks=["Document accessible to everyone about quality"],
            embeddings=[[0.1] * 1024],
            metadata={"title": "Open Doc"},
        )

        # Any user should see it
        results = knowledge_service.hybrid_search("quality", user_id=999)
        assert len(results) == 1


# ---------------------------------------------------------------------------
# Test: Celery Indexing Task (Task 12.6)
# ---------------------------------------------------------------------------


class TestIndexingTask:
    """Tests for the Celery indexing task configuration."""

    def test_task_retry_configuration(self) -> None:
        """index_document_task has exponential backoff retry configured."""
        from alcoabase.tasks.indexing_tasks import index_document_task

        # Verify retry configuration
        assert index_document_task.max_retries == 10
        assert index_document_task.autoretry_for is not None

    def test_task_processes_plain_text(
        self, knowledge_service: KnowledgeService
    ) -> None:
        """index_document_task processes plain text documents."""
        from alcoabase.tasks.indexing_tasks import (
            _get_knowledge_service,
            index_document_task,
        )

        # Patch the service getter to use our fixture
        with patch(
            "alcoabase.tasks.indexing_tasks._get_knowledge_service",
            return_value=knowledge_service,
        ):
            content = "This is test content for indexing. " * 20
            file_bytes_hex = content.encode("utf-8").hex()

            # Call the task function directly (not via Celery)
            result = index_document_task.apply(
                args=[
                    "2025-00001",
                    "1.0",
                    file_bytes_hex,
                    "text/plain",
                    {"title": "Test"},
                ]
            ).get()

            assert result["status"] == "indexed"
            assert result["chunk_count"] >= 1
