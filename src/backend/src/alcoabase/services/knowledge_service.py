"""Knowledge Service for document indexing, embedding, and hybrid search.

This module provides:
- Document text extraction (PDF, DOCX, plain text)
- Scanned PDF detection with OCR delegation
- Text chunking with configurable overlap
- Multilingual vector embedding generation (placeholder)
- OpenSearch indexing (placeholder)
- Hybrid search combining BM25 lexical + kNN semantic (placeholder)
- ABAC filtering and CSV record exclusion on search results

References:
    - Task 12: Knowledge Service (Document Indexing + Search)
    - Requirement 13: Document Indexing and Vector Embedding
    - Requirement 14: Semantic and Hybrid Search
"""

import logging
import random
from dataclasses import dataclass, field
from typing import Any

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from alcoabase.config import get_settings

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Data Classes
# ---------------------------------------------------------------------------


@dataclass
class SearchResult:
    """A single search result from hybrid search.

    Attributes:
        document_uuid: The Document-UUID of the matched document.
        title: Document title.
        version: Document version string (e.g., "1.0").
        excerpt: Matching text excerpt/chunk.
        relevance_score: Combined relevance score (0.0 to 1.0).
        metadata: Additional metadata (tags, language, etc.).
    """

    document_uuid: str
    title: str
    version: str
    excerpt: str
    relevance_score: float
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class IndexedDocument:
    """Represents a document stored in the in-memory index (placeholder).

    Attributes:
        document_uuid: The Document-UUID.
        version: Version string.
        chunks: Text chunks extracted from the document.
        embeddings: Vector embeddings for each chunk.
        metadata: Document metadata (title, tags, etc.).
        is_csv_validation_record: Whether this is a CSV validation record.
        permitted_user_ids: Set of user IDs with access (for ABAC filtering).
    """

    document_uuid: str
    version: str
    chunks: list[str]
    embeddings: list[list[float]]
    metadata: dict[str, Any] = field(default_factory=dict)
    is_csv_validation_record: bool = False
    permitted_user_ids: set[int] | None = None  # None means all users


# ---------------------------------------------------------------------------
# Knowledge Service
# ---------------------------------------------------------------------------


class KnowledgeService:
    """Service for document indexing, embedding generation, and hybrid search.

    Provides text extraction from multiple formats, chunking, embedding
    generation (placeholder), OpenSearch indexing (placeholder), and
    hybrid search with ABAC filtering.

    Attributes:
        _index: In-memory document index (placeholder for OpenSearch).
        _embedding_dimension: Dimension of embedding vectors from config.
    """

    def __init__(self) -> None:
        """Initialize KnowledgeService with settings and in-memory index."""
        settings = get_settings()
        self._embedding_dimension: int = settings.model_embedding_dimension
        self._index: dict[str, IndexedDocument] = {}

    # -----------------------------------------------------------------------
    # Text Extraction (Task 12.1)
    # -----------------------------------------------------------------------

    def extract_text(self, file_bytes: bytes, content_type: str) -> str:
        """Extract text content from a document file.

        Supports digital PDF (PyMuPDF), DOCX (python-docx), and plain text.

        Args:
            file_bytes: Raw file content as bytes.
            content_type: MIME type of the file (e.g., "application/pdf",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "text/plain").

        Returns:
            Extracted text content as a string.

        Raises:
            ValueError: If the content type is not supported.
        """
        if content_type == "application/pdf":
            return self._extract_pdf_text(file_bytes)
        elif content_type in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/docx",
        ):
            return self._extract_docx_text(file_bytes)
        elif content_type.startswith("text/"):
            return file_bytes.decode("utf-8", errors="replace")
        else:
            raise ValueError(f"Unsupported content type: {content_type}")

    def _extract_pdf_text(self, file_bytes: bytes) -> str:
        """Extract text from a digital PDF using PyMuPDF.

        Args:
            file_bytes: Raw PDF file content.

        Returns:
            Concatenated text from all PDF pages.
        """
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text_parts: list[str] = []
        for page in doc:
            page_text = page.get_text()
            if page_text.strip():
                text_parts.append(page_text)
        doc.close()
        return "\n".join(text_parts)

    def _extract_docx_text(self, file_bytes: bytes) -> str:
        """Extract text from a DOCX file using python-docx.

        Args:
            file_bytes: Raw DOCX file content.

        Returns:
            Concatenated text from all paragraphs.
        """
        import io

        doc = DocxDocument(io.BytesIO(file_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        return "\n".join(paragraphs)

    # -----------------------------------------------------------------------
    # Scanned PDF Detection (Task 12.2)
    # -----------------------------------------------------------------------

    def is_scanned_pdf(self, file_bytes: bytes) -> bool:
        """Detect if a PDF is scanned (image-based) with no extractable text.

        Checks each page for extractable text. If no pages contain text,
        the PDF is considered scanned.

        Args:
            file_bytes: Raw PDF file content.

        Returns:
            True if the PDF appears to be scanned (no extractable text).
        """
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        has_text = False
        for page in doc:
            page_text = page.get_text().strip()
            if page_text:
                has_text = True
                break
        doc.close()
        return not has_text

    def extract_text_with_ocr_fallback(
        self, file_bytes: bytes, content_type: str
    ) -> str:
        """Extract text with OCR fallback for scanned PDFs.

        If the PDF is scanned (no extractable text), delegates to OCR_Engine.
        Currently returns a placeholder message for scanned PDFs until
        the Model_Manager (Task 18) provides real OCR capability.

        Args:
            file_bytes: Raw file content.
            content_type: MIME type of the file.

        Returns:
            Extracted text content.
        """
        if content_type == "application/pdf" and self.is_scanned_pdf(file_bytes):
            logger.info("Scanned PDF detected, delegating to OCR_Engine (placeholder)")
            return self._ocr_extract_text(file_bytes)
        return self.extract_text(file_bytes, content_type)

    def _ocr_extract_text(self, file_bytes: bytes) -> str:
        """Placeholder for OCR text extraction via vision LLM.

        Will be implemented by Model_Manager (Task 18) with on-demand
        vision model loading.

        Args:
            file_bytes: Raw PDF file content.

        Returns:
            Placeholder text indicating OCR is pending.
        """
        logger.warning(
            "OCR_Engine not yet implemented. Returning placeholder for scanned PDF."
        )
        return "[OCR_PENDING: Scanned PDF text extraction requires Model_Manager]"

    # -----------------------------------------------------------------------
    # Text Chunking (Task 12.3)
    # -----------------------------------------------------------------------

    def chunk_text(
        self, text: str, chunk_size: int = 512, overlap: int = 50
    ) -> list[str]:
        """Split text into overlapping chunks for embedding generation.

        Uses whitespace-aware token approximation (splitting on whitespace).
        Each chunk contains approximately `chunk_size` tokens with `overlap`
        tokens of overlap between consecutive chunks.

        Args:
            text: Input text to chunk.
            chunk_size: Target number of tokens per chunk (default 512).
            overlap: Number of overlapping tokens between chunks (default 50).

        Returns:
            List of text chunks. Returns empty list for empty/whitespace input.

        Raises:
            ValueError: If overlap >= chunk_size.
        """
        if overlap >= chunk_size:
            raise ValueError(
                f"Overlap ({overlap}) must be less than chunk_size ({chunk_size})"
            )

        if not text or not text.strip():
            return []

        # Approximate tokens by splitting on whitespace
        tokens = text.split()

        if len(tokens) <= chunk_size:
            return [text.strip()]

        chunks: list[str] = []
        step = chunk_size - overlap
        i = 0

        while i < len(tokens):
            chunk_tokens = tokens[i : i + chunk_size]
            chunk_text = " ".join(chunk_tokens)
            chunks.append(chunk_text)

            if i + chunk_size >= len(tokens):
                break
            i += step

        return chunks

    # -----------------------------------------------------------------------
    # Embedding Generation (Task 12.4 - Placeholder)
    # -----------------------------------------------------------------------

    def generate_embeddings(self, chunks: list[str]) -> list[list[float]]:
        """Generate multilingual vector embeddings for text chunks.

        Placeholder implementation that returns random vectors of the
        correct dimension from config. The Model_Manager (Task 18) will
        provide real embedding generation via the multilingual-e5-large-instruct
        model.

        Args:
            chunks: List of text chunks to embed.

        Returns:
            List of embedding vectors, one per chunk. Each vector has
            dimension equal to `model_embedding_dimension` from settings.
        """
        embeddings: list[list[float]] = []
        for _ in chunks:
            # Generate a random unit vector of the correct dimension
            vec = [random.gauss(0, 1) for _ in range(self._embedding_dimension)]
            # Normalize to unit length
            magnitude = sum(v * v for v in vec) ** 0.5
            if magnitude > 0:
                vec = [v / magnitude for v in vec]
            embeddings.append(vec)
        return embeddings

    # -----------------------------------------------------------------------
    # OpenSearch Indexing (Task 12.5 - Placeholder)
    # -----------------------------------------------------------------------

    def index_document(
        self,
        document_uuid: str,
        version: str,
        chunks: list[str],
        embeddings: list[list[float]],
        metadata: dict[str, Any] | None = None,
    ) -> None:
        """Index document chunks and embeddings in OpenSearch.

        Placeholder implementation that stores in memory. Will be replaced
        with actual OpenSearch client calls when infrastructure is ready.

        Args:
            document_uuid: The Document-UUID to index.
            version: Document version string (e.g., "1.0").
            chunks: Text chunks extracted from the document.
            embeddings: Vector embeddings for each chunk.
            metadata: Optional metadata (title, tags, language, etc.).
        """
        if metadata is None:
            metadata = {}

        index_key = f"{document_uuid}:{version}"

        indexed_doc = IndexedDocument(
            document_uuid=document_uuid,
            version=version,
            chunks=chunks,
            embeddings=embeddings,
            metadata=metadata,
            is_csv_validation_record=metadata.get("is_csv_validation_record", False),
            permitted_user_ids=metadata.get("permitted_user_ids"),
        )

        self._index[index_key] = indexed_doc
        logger.info(
            "Indexed document %s v%s with %d chunks (placeholder)",
            document_uuid,
            version,
            len(chunks),
        )

    # -----------------------------------------------------------------------
    # Hybrid Search (Task 12.7 - Placeholder)
    # -----------------------------------------------------------------------

    def hybrid_search(
        self, query: str, user_id: int, limit: int = 20
    ) -> list[SearchResult]:
        """Perform hybrid search combining BM25 lexical + kNN semantic search.

        Placeholder implementation that performs simple keyword matching
        against the in-memory index. Will be replaced with actual OpenSearch
        hybrid query when infrastructure is ready.

        Args:
            query: Search query string.
            user_id: ID of the user performing the search (for ABAC filtering).
            limit: Maximum number of results to return (default 20).

        Returns:
            List of SearchResult objects ranked by relevance score.
        """
        query_lower = query.lower()
        results: list[SearchResult] = []

        for _key, doc in self._index.items():
            # Task 12.9: Exclude CSV validation records
            if doc.is_csv_validation_record:
                continue

            # Task 12.8: ABAC filtering
            if doc.permitted_user_ids is not None and user_id not in doc.permitted_user_ids:
                continue

            # Simple keyword matching (placeholder for BM25 + kNN)
            for i, chunk in enumerate(doc.chunks):
                if query_lower in chunk.lower():
                    # Calculate a simple relevance score
                    score = chunk.lower().count(query_lower) / max(len(chunk.split()), 1)
                    results.append(
                        SearchResult(
                            document_uuid=doc.document_uuid,
                            title=doc.metadata.get("title", "Untitled"),
                            version=doc.version,
                            excerpt=chunk[:200],
                            relevance_score=min(score, 1.0),
                            metadata=doc.metadata,
                        )
                    )

        # Sort by relevance score descending
        results.sort(key=lambda r: r.relevance_score, reverse=True)
        return results[:limit]
